#!/usr/bin/env python3
import sys, json, copy, subprocess, shutil, os, tempfile
from docx import Document
from lxml import etree

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
XML_SPACE = '{http://www.w3.org/XML/1998/namespace}space'

def all_text(el):
    return ''.join((t.text or '') for t in el.iter(f'{{{W}}}t'))

def clear_runs(cell):
    for p in cell.paragraphs:
        for r in p.runs: r.text = ''

def set_cell(cell, texto, sz=14):
    from docx.oxml.ns import qn
    for p in cell.paragraphs:
        runs = p.runs
        if runs:
            runs[0].text = texto
            for r in runs[1:]: r.text = ''
            return
    # No runs — crear uno con tamaño correcto
    run = cell.paragraphs[0].add_run(texto)
    rpr = etree.SubElement(run._element, f'{{{W}}}rPr')
    run._element.insert(0, rpr)
    for tag in ['sz', 'szCs']:
        el = etree.SubElement(rpr, f'{{{W}}}'+tag)
        el.set(f'{{{W}}}val', str(sz))

def copiar_fila_despues(tabla, idx):
    tr = tabla.rows[idx]._tr
    tr_nuevo = copy.deepcopy(tr)
    tr.addnext(tr_nuevo)

def set_para_total(child, texto):
    """Escribe el total en un párrafo existente."""
    for t_el in child.iter(f'{{{W}}}t'): t_el.text = ''
    runs_el = child.findall(f'.//{{{W}}}r')
    if runs_el:
        t_els = runs_el[0].findall(f'{{{W}}}t')
        if t_els:
            t_els[0].text = texto
            t_els[0].set(XML_SPACE, 'preserve')
        else:
            new_t = etree.SubElement(runs_el[0], f'{{{W}}}t')
            new_t.text = texto
            new_t.set(XML_SPACE, 'preserve')
    else:
        r_el = etree.SubElement(child, f'{{{W}}}r')
        rpr = etree.SubElement(r_el, f'{{{W}}}rPr')
        etree.SubElement(rpr, f'{{{W}}}b')
        c_el = etree.SubElement(rpr, f'{{{W}}}color')
        c_el.set(f'{{{W}}}val', '00B050')
        sz_el = etree.SubElement(rpr, f'{{{W}}}sz')
        sz_el.set(f'{{{W}}}val', '22')
        t_el = etree.SubElement(r_el, f'{{{W}}}t')
        t_el.text = texto
        t_el.set(XML_SPACE, 'preserve')

def generar(datos):
    doc      = Document('/app/Cotizacion Qurakuna-PLANTILLA.docx')
    nombre   = datos['nombre']
    productos= datos.get('productos', [])
    traslado = datos.get('traslado', {})
    inc_tras = traslado.get('incluye', False)
    p_tras   = float(traslado.get('precio', 0))
    d_tras   = traslado.get('descripcion', 'Recojo de Plantas & Traslado de Plantas, Personal, Macetas. Ida y Vuelta.')
    pack     = datos.get('pack_cuidado', False)
    base     = datos.get('base_movil', False)
    mano     = datos.get('mano_obra', {})
    inc_mano = mano.get('incluye', False)
    p_mano   = float(mano.get('precio', 0))
    d_mano   = mano.get('descripcion', 'Servicio de instalación por jardinero.')

    # ── 1. Nombre en párrafos ──
    for p in doc.paragraphs:
        for r in p.runs:
            if '[NOMBRE DEL CLIENTE]' in (r.text or ''):
                r.text = r.text.replace('[NOMBRE DEL CLIENTE]', nombre)

    # ── 2. Tabla de productos ──
    t0 = doc.tables[0]
    n_prod = len(productos)
    while len(t0.rows) - 1 < max(n_prod, 1):
        copiar_fila_despues(t0, len(t0.rows) - 1)
    for i, prod in enumerate(productos):
        fila   = t0.rows[i + 1]
        precio = float(prod.get('precio', 0))
        cant   = int(prod.get('cantidad', 1))
        set_cell(fila.cells[0], prod.get('espacio', ''))
        set_cell(fila.cells[1], prod.get('nombre', ''))
        set_cell(fila.cells[2], prod.get('descripcion', ''))
        set_cell(fila.cells[3], f'S/.{precio:,.2f}')
        set_cell(fila.cells[4], str(cant).zfill(2))
        set_cell(fila.cells[5], f'S/.{precio*cant:,.2f}')
    # Eliminar filas sobrantes
    for i in range(len(t0.rows) - 1, n_prod, -1):
        tr = t0.rows[i]._tr
        tr.getparent().remove(tr)

    # ── 3. Total ──
    total = sum(float(p.get('precio',0)) * int(p.get('cantidad',1)) for p in productos)
    if inc_tras: total += p_tras
    if inc_mano: total += p_mano
    if pack:     total += 60
    if base:     total += 60

    # ── 4. Traslado ──
    t1 = doc.tables[1]
    body_ch = list(doc.element.body)
    if inc_tras and p_tras > 0:
        fila_t = t1.rows[1]
        # descripción
        p_d = fila_t.cells[1].paragraphs[0]
        for r in p_d.runs: r.text = ''
        if p_d.runs: p_d.runs[0].text = d_tras
        else: p_d.add_run(d_tras)
        # precio y costo
        for col in [2, 4]:
            target_p = next((p for p in fila_t.cells[col].paragraphs if p.runs), fila_t.cells[col].paragraphs[0])
            txt = f'S/. {p_tras:,.2f}'
            if target_p.runs:
                target_p.runs[0].text = txt
                for r in target_p.runs[1:]: r.text = ''
            else: target_p.add_run(txt)
    # Mano de obra (fila 2 de tabla traslado, si existe)
    if inc_mano and p_mano > 0:
        # Si no hay fila 2, copiarla de la fila 1
        if len(t1.rows) < 3:
            copiar_fila_despues(t1, 1)
        fila_m = t1.rows[2]
        set_cell(fila_m.cells[0], 'Mano de obra')
        set_cell(fila_m.cells[1], d_mano)
        set_cell(fila_m.cells[2], f'S/. {p_mano:,.2f}')
        set_cell(fila_m.cells[3], '1.00')
        set_cell(fila_m.cells[4], f'S/. {p_mano:,.2f}')
    else:
        # Vaciar fila de mano de obra si existe
        if len(t1.rows) > 2:
            for cell in t1.rows[2].cells: clear_runs(cell)

    if not inc_tras and not inc_mano:
        # Eliminar título y tabla completa
        to_remove = []
        for child in list(doc.element.body):
            txt = all_text(child)
            tag = child.tag.split('}')[1] if '}' in child.tag else child.tag
            if 'Servicio Adicional' in txt:
                to_remove.append(child)
            elif tag == 'tbl' and 'Traslado' in txt and 'Servicio' in txt:
                to_remove.append(child)
        for child in to_remove:
            child.getparent().remove(child)

    # ── 5. Opcionales ──
    t2 = doc.tables[2]
    if not pack and not base:
        # Eliminar bloque de opcionales incluyendo línea divisoria y párrafos vacíos previos
        body_now = list(doc.element.body)
        to_remove = []
        in_op = False
        # También eliminar el separador (---) y párrafos vacíos antes de "Productos Opcionales"
        for i, child in enumerate(body_now):
            txt = all_text(child)
            tag = child.tag.split('}')[1] if '}' in child.tag else child.tag
            if '---' in txt or (tag == 'p' and not txt.strip() and in_op is False):
                # Check if next non-empty element is "Productos Opcionales"
                for j in range(i+1, len(body_now)):
                    next_txt = all_text(body_now[j]).strip()
                    if next_txt:
                        if 'Productos Opcionales' in next_txt:
                            to_remove.append(child)
                        break
            if 'Productos Opcionales' in txt:
                in_op = True
                to_remove.append(child)
                continue
            if in_op:
                if 'Total' in txt or 'Precios no incluyen' in txt:
                    break
                to_remove.append(child)
        for child in to_remove:
            try: child.getparent().remove(child)
            except: pass
    else:
        if not pack:
            for cell in t2.rows[1].cells: clear_runs(cell)
        if not base:
            for cell in t2.rows[2].cells: clear_runs(cell)

    # ── 6. Escribir Total en párrafo [18] = "  Total = " ──
    total_txt = f'Total: S/.{total:,.2f}'
    for child in list(doc.element.body):
        txt = all_text(child)
        if 'Total' in txt and '=' in txt:
            set_para_total(child, total_txt)
            break

    # ── 7. Guardar y convertir ──
    tmpdir = tempfile.mkdtemp()
    safe   = nombre.replace(' ','_').replace('/','_')
    docx_p = os.path.join(tmpdir, f'Cotizacion_Qurakuna_{safe}.docx')
    doc.save(docx_p)

    subprocess.run(
        ['libreoffice','--headless','--convert-to','pdf','--outdir',tmpdir,docx_p],
        capture_output=True, text=True, timeout=60
    )

    pdf_p = docx_p.replace('.docx', '.pdf')
    os.makedirs('/tmp/cotizaciones', exist_ok=True)
    out   = f'/tmp/cotizaciones/Cotizacion_Qurakuna_{safe}.pdf'
    if os.path.exists(pdf_p):
        shutil.copy(pdf_p, out)
        print(out)
    else:
        out_d = out.replace('.pdf','.docx')
        shutil.copy(docx_p, out_d)
        print(out_d)

    shutil.rmtree(tmpdir, ignore_errors=True)

if __name__ == '__main__':
    datos = json.loads(sys.argv[1])
    generar(datos)
